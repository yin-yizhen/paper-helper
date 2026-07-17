#!/usr/bin/env python3
"""Build a clean, basic Chinese course-paper DOCX from Markdown and references.json."""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path


def parse_markdown(source: str) -> tuple[str, list[tuple[str, int, str]]]:
    title = "课程论文"
    blocks: list[tuple[str, int, str]] = []
    for raw in source.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# ") and title == "课程论文":
            title = line[2:].strip()
        elif match := re.match(r"^(#{1,6})\s+(.+)$", line):
            blocks.append(("heading", len(match.group(1)), match.group(2).strip()))
        else:
            blocks.append(("paragraph", 0, line))
    return title, blocks


def reference_text(item: dict) -> str:
    authors = "，".join(item.get("authors", []))
    vi = str(item.get("volume") or "")
    if item.get("issue"):
        vi += f"({item['issue']})"
    pages = f":{item['pages']}" if item.get("pages") else ""
    return f"[{item['id']}] {authors}. {item['title']}[J]. {item['journal']}, {item['year']}{', ' + vi if vi else ''}{pages}."


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("manuscript", type=Path)
    parser.add_argument("--references", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError:
        print("缺少 python-docx；请在具备文档生成能力的 Codex 环境运行。", file=sys.stderr)
        return 2
    try:
        title, blocks = parse_markdown(args.manuscript.read_text(encoding="utf-8"))
        refs = json.loads(args.references.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        print(f"无法读取输入：{exc}", file=sys.stderr)
        return 2
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.0)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    for kind, level, text in blocks:
        if kind == "heading":
            doc.add_heading(text, level=min(level, 3))
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
    doc.add_heading("参考文献", level=1)
    for item in refs:
        p = doc.add_paragraph(reference_text(item))
        p.paragraph_format.hanging_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"已生成 {args.output}。请按 documents 工作流渲染并逐页检查。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
